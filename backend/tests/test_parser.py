from pathlib import Path

import pytest

from backend.app.core.parser import extract_text_from_file


def test_extract_text_from_txt_file(tmp_path: Path):
    resume = tmp_path / "resume.txt"
    resume.write_text("Python developer with SQL experience.", encoding="utf-8")

    assert extract_text_from_file(resume) == "Python developer with SQL experience."


def test_extract_text_from_docx_file(tmp_path: Path):
    from docx import Document

    resume = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Data analyst with Python and Excel.")
    document.add_paragraph("Built dashboards using Power BI.")
    document.save(resume)

    text = extract_text_from_file(resume)

    assert "Data analyst" in text
    assert "Power BI" in text


def test_extract_text_rejects_unsupported_file(tmp_path: Path):
    resume = tmp_path / "resume.png"
    resume.write_text("not a resume", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text_from_file(resume)


def test_extract_text_reports_invalid_docx(tmp_path: Path):
    resume = tmp_path / "broken.docx"
    resume.write_text("not actually a docx", encoding="utf-8")

    with pytest.raises(ValueError, match="Could not read DOCX text"):
        extract_text_from_file(resume)
